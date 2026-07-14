from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("AI-Word文档导入测试样例.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)


def set_cell_fill(cell, color: str):
    """中文注解：直接设置单元格底色，确保 Word 与 LibreOffice 渲染一致。"""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def configure_styles(document: Document):
    """中文注解：显式配置正文和标题样式，用于测试导入后层级与格式识别。"""
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in (
        ("Title", 24, DARK),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
    ):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def build_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI Word 文档助手项目周报")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("导入与在线编辑功能测试样例 | 2026 年 7 月").italic = True

    document.add_heading("一、本周工作概览", level=1)
    document.add_paragraph(
        "本周围绕 AI Word 文档助手的文档导入能力开展开发与验证，重点完成 DOCX 解析、"
        "文字型 PDF 提取、MySQL 文档入库以及 MinIO 原文件归档。用户导入文件后，系统会自动创建文档记录并打开编辑器。"
    )
    document.add_paragraph(
        "该能力让用户可以将已有材料继续交给 AI 润色、扩写和格式优化，减少重复复制粘贴，形成从存量文档到智能编辑的完整流程。"
    )

    document.add_heading("二、已完成功能", level=1)
    for item in (
        "支持上传最大 15MB 的 DOCX 文档。",
        "支持提取文字型 PDF 的正文内容。",
        "导入后自动加入最近文档列表并打开编辑器。",
        "原始文件归档至 MinIO，文档内容和索引保存至 MySQL。",
        "导入内容经过白名单清洗，避免不安全标签进入编辑器。",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("三、测试操作步骤", level=1)
    for item in (
        "在工作台点击“导入文档”。",
        "选择本测试样例文件并等待导入完成。",
        "确认页面自动进入正文编辑器。",
        "修改任意段落并点击保存。",
        "使用 AI 润色或格式优化功能处理选中文字。",
        "导出 Word，检查内容和样式是否符合预期。",
    ):
        document.add_paragraph(item, style="List Number")

    document.add_heading("四、验收记录", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = (Inches(2.0), Inches(3.4), Inches(1.1))
    headers = ("检查项", "预期结果", "状态")
    for index, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, "F2F4F7")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    for row_values in (
        ("标题识别", "一级标题出现在文档大纲中", "待测试"),
        ("正文编辑", "段落可以选择、修改并保存", "待测试"),
        ("列表格式", "项目符号和编号内容可正常阅读", "待测试"),
        ("AI 编辑", "选中文字可以润色、扩写或缩写", "待测试"),
        ("Word 导出", "导出文件可以正常打开", "待测试"),
    ):
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row_values)):
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = text

    document.add_heading("五、测试说明", level=1)
    note = document.add_paragraph()
    lead = note.add_run("注意：")
    lead.bold = True
    lead.font.color.rgb = DARK
    note.add_run(
        "当前版本主要还原文字、标题和列表。文档中的复杂表格、图片、页眉页脚以及浮动对象，"
        "导入编辑器时可能转换为普通内容或提示占位，这是第一版预期行为。扫描版 PDF 需要 OCR 能力后才能识别。"
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
